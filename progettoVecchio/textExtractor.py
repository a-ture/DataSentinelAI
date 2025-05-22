import io
import PyPDF2
import docx
import re
from fpdf import FPDF

def extract(file):
    text = ""
    extension = findExtension(file)
    uploaded_file = io.BytesIO(file.read())
    match extension:
        case ".docx":
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        case ".pdf":
            pdf = PyPDF2.PdfFileReader(uploaded_file)
            for page in range(pdf.getNumPages()):
                text += pdf.getPage(page).extract_text()
            return text
        case ".txt":
            text = uploaded_file.read()
            return text
    return None

def findExtension(file):
    return re.findall(r"\.+[A-Za-z]*", file.name).pop()

def writeFile(data, extension):
    match extension:
        case ".docx":
            doc = docx.Document()
            doc.add_paragraph(data)
            doc.save("edited.docx")
            return open("edited.docx", "rb").read()
        case ".pdf":
            pdf=FPDF()
            pdf.core_fonts_encoding = 'utf8'
            pdf.add_page(orientation="P")
            pdf.set_font("helvetica")
            pdf.multi_cell(w=0, h=10, text=data, border=0, align="L")
            pdf.output("edited.pdf", "S")
            return open("edited.pdf", "rb").read()
        case ".txt":
            return data
    return None
